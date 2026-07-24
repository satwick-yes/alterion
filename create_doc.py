import docx
import os

doc = docx.Document()

doc.add_heading('Potential IEEE Conference Questions', 0)

doc.add_paragraph('Paper: Mathematical Modelling of a Dual-Use Electromagnetic Coil for Simultaneous Magnetic Levitation and Wireless Power Transfer in Intelligent Transportation Systems')
doc.add_paragraph('Authors: Kavita Jindal, Satwick Shaw\n')

doc.add_heading('1. Novelty & Core Contributions', level=1)
doc.add_paragraph('Q1: How does your time-frequency multiplexing approach differ from existing dual-mode electromagnetic designs, and what is the primary advantage of combining them into a single coil?')
doc.add_paragraph('Q2: You mentioned that traditional systems use separate coils because of differing frequencies. How did you completely isolate and overcome the interference between the low-frequency levitation signal and the high-frequency WPT signal within the same structure?')

doc.add_heading('2. Limitations & Stability', level=1)
doc.add_paragraph('Q3: The mathematical model assumes certain ideal conditions. How do non-linearities in the magnetic core or lateral/vertical misalignments of the vehicle affect the stability of the levitation and the efficiency of the WPT?')
doc.add_paragraph('Q4: Your thermal constraint analysis highlights copper and core losses. In a real-world continuous operation scenario, would passive cooling be sufficient, or does this dual-use coil require active cooling mechanisms to maintain ΔT ≤ ΔT_max?')

doc.add_heading('3. Future Work & Physical Prototyping', level=1)
doc.add_paragraph('Q5: The current findings are based on mathematical modelling and simulations. What are the main physical and material challenges you anticipate when moving to a hardware prototype?')
doc.add_paragraph('Q6: How scalable is this single-coil architecture for heavier transport vehicles like full-sized MagLev trains, as opposed to lighter automated transit pods?')

doc.add_heading('4. Practical Implications', level=1)
doc.add_paragraph('Q7: Implementing this dual-use coil in intelligent transportation systems would require significant infrastructure changes. How do the installation and maintenance costs realistically compare to existing separated MagLev and WPT systems?')
doc.add_paragraph('Q8: Dynamic charging involves moving receivers. How sensitive is the wireless power transfer efficiency to the high-speed motion of the vehicle over the guideway coil, considering the skin and proximity effects you mentioned?')

desktop = os.path.join(os.path.expanduser("~"), "Desktop")
path = os.path.join(desktop, "potential_conference_questions.docx")
doc.save(path)
print(f"Document saved to {path}")
